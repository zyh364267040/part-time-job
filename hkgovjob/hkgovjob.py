# -*- coding = utf-8 -*-
# @Time: 2022/10/3 20:06
import requests
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm

from PASSWORD import *


# 创建word文档
doc = Document()
# 设置字体
doc.styles['Normal'].font.name = u'宋体'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 将段落中的所有字体
# doc.styles["Normal"].font.name = u'宋体'


# post请求登录
def post_data():
    # 创建session
    session = requests.session()
    url = 'http://hkgovjob.com/?page_id=2351'
    data = {
        'username': HKGOVJOB_USERNAME,
        'password': HKGOVJOB_PASSWORD,
        '_n': 'b16bf5025f',
        '_wp_http_referer': '/?page_id=2351',
        'redirect': '',
        'login': '登錄'
    }
    headers = {
        'user-agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/104.0.0.0 Safari/537.36'
    }
    session.post(url, headers=headers, data=data)
    return session


def get_data(session, page):
    headers = {
        'user-agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/104.0.0.0 Safari/537.36'
    }
    # 中文试题
    # url = f'http://hkgovjob.com/?page_id=3229&paged={page}'
    # url = 'http://hkgovjob.com/?page_id=3229&paged=10'
    # 英文试题
    url = f'http://hkgovjob.com/?page_id=3237&paged={page}'
    res = session.get(url, headers=headers)
    return res


def parse_data(res):
    # 使用xpath解析数据
    tree = etree.HTML(res.text)
    article_list = tree.xpath('//div[@class="template-blog"]/article')
    # print(1, article_list)
    return article_list


def save_data(article_list):
    # 循环每一个试题
    for article in article_list:
        title = article.xpath('./div/header/h2/a/text()')[0].replace('\t', '')
        doc.add_heading(title)
        print(title)

        info_list = article.xpath('./div/header/span//text()')
        info = ''.join(info_list)
        doc.add_paragraph(info)

        trs = article.xpath('./div[2]/div[1]/table/tbody/tr')
        # print(2, trs)
        if not trs:
            text_list = article.xpath('./div[2]/div[1]//text()')
            text = ''.join(text_list)
            # print(text)
            doc.add_paragraph(text)
            doc.add_page_break()
            continue

        # 创建表格
        table = doc.add_table(rows=0, cols=2)

        for tr in trs:
            td1 = tr.xpath('./td[1]//text()')
            if '\n' in td1:
                for i in range(len(td1)-1, -1, -1):
                    if td1[i] == '\n':
                        td1.pop(i)
            td2 = tr.xpath('./td[2]//text()')
            if '\n' in td2:
                for i in range(len(td2)-1, -1, -1):
                    if td2[i] == '\n':
                        td2.pop(i)

            if not td2:
                print(td1)
                row_cells = table.add_row().cells
                if len(td1) > 1:
                    td1 = ''.join(td1)
                else:
                    if not td1:
                        continue
                    td1 = td1[0]
                row_cells[0].text = td1
            else:
                row_cells = table.add_row().cells
                if len(td1) > 1:
                    td1 = ''.join(td1)
                else:
                    if not td1:
                        continue
                    td1 = td1[0]

                if len(td2) > 1:
                    td2 = ''.join(td2)
                else:
                    td2 = td2[0]
                row_cells[0].text = td1
                row_cells[1].text = td2
        # 切换到下一页
        doc.add_page_break()


def main():
    # 发送post请求
    session = post_data()
    # for page in range(1, 39):
    for page in range(1, 34):
        print(f'开始获取第{page}页数据...')

        # 发送get请求,获取数据
        res = get_data(session, page)
        print(f'第{page}页数据获取成功!!!')

        # 解析数据
        article_list = parse_data(res)

        # 保存数据
        save_data(article_list)
        print(f'第{page}页数据保存成功!!!')
        # break

    # 保存word文档
    # doc.save('熱門中文試題.docx')
    doc.save('熱門英文試題.docx')
    print('所有数据获取成功!!!')


if __name__ == '__main__':
    main()
